#!/usr/bin/env python3
"""Generate Innovation Depot application Word document from the updated markdown draft."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# Mealvana brand colors
BLACKBERRY = RGBColor(0x3D, 0x1F, 0x47)
ORANGE = RGBColor(0xFF, 0x8B, 0x3D)
TEXT_DARK = RGBColor(0x38, 0x16, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
DRAGONFRUIT = RGBColor(0xE8, 0x43, 0x93)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = TEXT_DARK

for level, (size, color) in {1: (22, BLACKBERRY), 2: (16, BLACKBERRY), 3: (12, ORANGE)}.items():
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(18 if level < 3 else 12)
    s.paragraph_format.space_after = Pt(4)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3D1F47')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_qa(doc, question, answer, required=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(question)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    if required:
        run2 = p.add_run(' *')
        run2.font.color.rgb = DRAGONFRUIT
        run2.bold = True

    p = doc.add_paragraph(answer)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


# Read the updated markdown
with open(os.path.expanduser('~/development/mealvana_endurance/docs/_archived/business/innovation_depot_application_draft.md'), 'r') as f:
    md = f.read()

# Parse sections and Q&A from markdown
sections = md.split('\n---\n')

# Helper to extract Q&A pairs from a section
def extract_qa(text):
    pairs = []
    parts = re.split(r'\n### ', text)
    for part in parts[1:]:  # skip header
        lines = part.strip().split('\n', 1)
        question = lines[0].strip()
        answer = lines[1].strip() if len(lines) > 1 else ''
        # Clean markdown formatting from answer
        answer = re.sub(r'\*\*([^*]+)\*\*', r'\1', answer)  # remove bold
        answer = answer.replace('- **', '- ').replace('**', '')
        pairs.append((question, answer))
    return pairs

# ========== TITLE PAGE ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('INNOVATION DEPOT')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Member Resource and\nSupport Application')
run.font.size = Pt(28)
run.font.color.rgb = BLACKBERRY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('Milkman, Inc. (Mealvana Endurance)')
run.font.size = Pt(16)
run.font.color.rgb = ORANGE
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('April 2026')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('Prepared by Xuan Huang & Lee Martin')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ========== Parse and render each section ==========
section_map = {
    'COMPANY FOUNDATIONS': ('Company Foundations', 'Baseline information about your company and team'),
    'SUPPORT REQUEST': ('Support Request', 'What you are trying to achieve, what is getting in the way, and which resources would help most'),
    'SUPPORT PRIORITIES': ('Support Priorities', 'Rate how much you need each type of support right now'),
    'BUSINESS & MARKET': ('Business & Market', 'The opportunity you are building into and the strength of your model'),
    'TEAM & FOUNDER': ('Team & Founder', 'Who you are, what makes you uniquely positioned, and how you lead'),
    'PRODUCT & BRAND': ('Product & Brand', 'Where your product and brand stand today \u2014 and where they are headed'),
}

for section_text in sections:
    # Find which section this is
    matched_key = None
    for key in section_map:
        if key in section_text:
            matched_key = key
            break

    if not matched_key:
        continue

    title, subtitle = section_map[matched_key]

    doc.add_heading(title, level=1)
    p = doc.add_paragraph(subtitle)
    p.runs[0].font.color.rgb = GRAY
    p.runs[0].font.italic = True
    add_divider(doc)

    # Extract Q&A pairs
    pairs = extract_qa(section_text)
    for question, answer in pairs:
        # Determine if optional
        required = True
        if 'Regulatory' in question or 'Intellectual Property' in question or 'Additional Materials' in question:
            required = False
        add_qa(doc, question, answer, required=required)

    doc.add_page_break()

# Remove the notes section by not including it
# The markdown has a NOTES FOR REVIEW section we skip

output_path = os.path.expanduser('~/development/mealvana_endurance/docs/_archived/business/Innovation_Depot_Application_Mealvana.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
