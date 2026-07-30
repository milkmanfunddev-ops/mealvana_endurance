#!/usr/bin/env python3
"""Generate Innovation Depot application Word document with exact questions and answers."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Mealvana brand colors
BLACKBERRY = RGBColor(0x3D, 0x1F, 0x47)
ORANGE = RGBColor(0xFF, 0x8B, 0x3D)
CREAM_BG = RGBColor(0xF5, 0xF3, 0xED)
ELECTROLYTE = RGBColor(0x5D, 0xE4, 0xD3)
TEXT_DARK = RGBColor(0x38, 0x16, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = TEXT_DARK

style = doc.styles['Heading 1']
font = style.font
font.name = 'Calibri'
font.size = Pt(22)
font.color.rgb = BLACKBERRY
font.bold = True
style.paragraph_format.space_before = Pt(24)
style.paragraph_format.space_after = Pt(6)

style = doc.styles['Heading 2']
font = style.font
font.name = 'Calibri'
font.size = Pt(16)
font.color.rgb = BLACKBERRY
font.bold = True
style.paragraph_format.space_before = Pt(18)
style.paragraph_format.space_after = Pt(4)

style = doc.styles['Heading 3']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = ORANGE
font.bold = True
style.paragraph_format.space_before = Pt(12)
style.paragraph_format.space_after = Pt(4)


def add_divider(doc):
    """Add a subtle horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3D1F47')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_qa(doc, question, answer, required=True):
    """Add a question-answer pair with clear formatting."""
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(question)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    if required:
        run2 = p.add_run(' *')
        run2.font.color.rgb = RGBColor(0xE8, 0x43, 0x93)  # dragonfruit
        run2.bold = True

    # Answer
    if isinstance(answer, list):
        for line in answer:
            p = doc.add_paragraph(line, style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = TEXT_DARK
    else:
        p = doc.add_paragraph(answer)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_DARK


# ========== TITLE PAGE ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('INNOVATION DEPOT')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'
run.bold = False

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Member Resource and\nSupport Application')
run.font.size = Pt(28)
run.font.color.rgb = BLACKBERRY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('Milkman, Inc. (Mealvana Endurance)')
run.font.size = Pt(16)
run.font.color.rgb = ORANGE
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('April 2026')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('Prepared by Xuan Huang & Lee Martin')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

doc.add_page_break()

# ========== SECTION 1: COMPANY FOUNDATIONS ==========
doc.add_heading('Company Foundations', level=1)
p = doc.add_paragraph('Baseline information about your company and team')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'Company Name', 'Milkman, Inc.')

add_qa(doc, 'Founder Email', 'xh.analytics@gmail.com, lee.b.martin@gmail.com')

add_qa(doc, 'Year Founded', '2019')

add_qa(doc, 'Website & Social Presence',
       'Please share your company website and primary social media profiles.')
# Replace with detailed answer
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
lines = [
    ('Product website: ', 'https://endurance.mealvana.io/'),
    ('Company website: ', 'https://www.mealvana.io/'),
    ('LinkedIn (Xuan Huang, Founder): ', 'https://www.linkedin.com/in/xuan-huang/'),
    ('LinkedIn (Lee Martin, CTO): ', 'https://www.linkedin.com/in/lee-martin-developer'),
    ('Instagram: ', 'https://www.instagram.com/mealvana_endurance/'),
    ('App Store (iOS): ', 'https://apps.apple.com/us/app/mealvana-endurance/id6751113738'),
    ('Google Play: ', 'https://play.google.com/store/apps/details?id=com.milkman.mealvanaendurance&pli=1'),
]
for i, (label, value) in enumerate(lines):
    if i > 0:
        p.add_run('\n')
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK

add_qa(doc, 'Industry / Sector', 'Sports Nutrition Technology / Health & Fitness')

add_qa(doc, 'Company Stage', 'Pre-revenue')

add_qa(doc, 'Full-Time Team Size', '2')

add_qa(doc, 'Prior Funding',
       'We have not raised outside equity capital. Our external funding includes:\n\n'
       '\u2022 $25,000 Alabama Launchpad prize (November 2021) \u2014 Non-dilutive concept-stage prize from Alabama Launchpad Cycle 3 Finale, administered by the Economic Development Partnership of Alabama.\n'
       '\u2022 Innovation Depot Special Projects Grant \u2014 Non-dilutive grant used to fund Dr. Rachel Mitchell (sports dietitian) for clinical validation of our nutrition algorithms and development of our Mealvana 101 educational content. [CONFIRM: amount and date]\n'
       '\u2022 Consulting revenue from Piggly Wiggly \u2014 [CONFIRM: amount, date, and scope of work]\n\n'
       'There are no outside equity investors, active advisors from prior funding, or board members. The cap table is clean \u2014 100% founder-held. The company has been bootstrapped by the two co-founders since inception, supplemented by the non-dilutive grants and consulting revenue above.')

doc.add_page_break()

# ========== SECTION 2: SUPPORT REQUEST ==========
doc.add_heading('Support Request', level=1)
p = doc.add_paragraph('What you are trying to achieve, what is getting in the way, and which resources would help most')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'Business Goals \u2014 Next 12 Months',
       'Our top three business goals for the next 12 months are:\n\n'
       '1. Launch paid subscription and reach 100+ paying subscribers by Q1 2027. We have built the full product \u2014 AI-powered nutrition planning, coach mode, integrations with TrainingPeaks, FinalSurge, and Garmin \u2014 and are now implementing RevenueCat to gate 11 identified Pro features (training platform integrations, coach dashboard, brick workout planning, carb-loading protocols, personal templates, barcode scanning, and more). Our pricing target is $9.99/month or $69.99/year with a 1-month free trial.\n\n'
       '2. Grow to 1,000+ app downloads and onboard 10+ coaches to Coach Mode. Coach Mode v2 is production-ready with many-to-many coach-athlete relationships, email invitations, permission levels, messaging, and a web-based dashboard. We need to drive initial adoption among RRCA-certified running coaches and sports dietitians to validate the B2B revenue channel (targeting $20\u2013$50/month per coach subscription).\n\n'
       '3. Close a pre-seed round of $250K\u2013$500K within 3\u20136 months to fund a go-to-market hire, extend runway, and accelerate user acquisition. This fundraise will be our first outside capital beyond the Alabama Launchpad prize.')

add_qa(doc, 'Core Problem / Constraint',
       'Capital and runway. We are a two-person founding team that has bootstrapped a production-grade, feature-rich product \u2014 but we are resource-constrained on two fronts that capital would unlock:\n\n'
       'First, we lack dedicated go-to-market capacity. Both founders are spending the majority of their time on product development, which means user acquisition, content marketing, coach outreach, and partnership development receive insufficient attention. We need either a dedicated hire or funded marketing initiatives to drive the top of the funnel.\n\n'
       'Second, our runway limits how aggressively we can pursue the revenue milestones that would make us investable. We have a live product with 100\u2013500 users, positive beta feedback, coach interest, and a clear monetization path \u2014 but converting that into paying subscribers and closing enterprise pilot deals requires sustained effort that our current bandwidth cannot support.\n\n'
       'The constraint is not product or technology \u2014 it is the resources to bring what we have already built to market effectively.')

add_qa(doc, 'How You Will Measure Success',
       'Over the next 12 months, we will track:\n\n'
       '\u2022 App installs: Target 1,000+ total downloads across iOS and Android (currently 100\u2013500 users)\n'
       '\u2022 Paid conversion: Target 100+ paying subscribers at $9.99/month or $69.99/year, representing ~$12K\u2013$70K ARR depending on mix\n'
       '\u2022 Coach Mode adoption: Target 10+ verified coaches onboarded, each managing 5\u201310 athletes, as validation for the B2B revenue channel\n'
       '\u2022 Fundraise: Close a $250K\u2013$500K pre-seed round within 3\u20136 months\n'
       '\u2022 Retention: Establish baseline retention metrics via Mixpanel (we have analytics infrastructure in place but are still building a meaningful data set)\n'
       '\u2022 Coach Mode MRR: Reach $1K\u2013$2K MRR from coach subscriptions within 6 months of launch\n\n'
       '\u201cWinning\u201d by Q1 2027 means: subscription revenue live, 100+ paying users demonstrating willingness to pay, a closed pre-seed round, and 10+ coaches actively using the platform \u2014 proving both B2C and B2B product-market fit.')

add_qa(doc, 'Works in Progress',
       'Yes, we are actively working on the capital constraint:\n\n'
       'What we have tried:\n'
       '\u2022 Successfully applied for and received an Innovation Depot Special Projects Grant, which funded Dr. Rachel Mitchell (sports dietitian) for clinical validation of our nutrition algorithms \u2014 this directly improved the scientific rigor of our product\n'
       '\u2022 Generated consulting revenue through a project with Piggly Wiggly, demonstrating our team\u2019s ability to deliver value in adjacent nutrition/data spaces while building Mealvana\n'
       '\u2022 Won the $25K Alabama Launchpad prize (2021) for concept validation\n'
       '\u2022 Maintained extremely lean operations \u2014 two full-time co-founders, supplemented by a part-time user researcher (Rui Tian), a part-time sports dietitian (Dr. Rachel Mitchell), and an engineering consultant (Haibin Zhuang)\n'
       '\u2022 Continued aggressive bootstrapping: shipped 18 app versions, built Coach Mode, completed integrations with TrainingPeaks, FinalSurge, and Garmin, and built a web portal \u2014 all on personal capital\n\n'
       'What we learned:\n'
       '\u2022 The product is strong enough to attract organic users and positive feedback, but without dedicated marketing and sales bandwidth, growth is linear rather than exponential\n'
       '\u2022 Coach interest is real but requires active outreach and relationship-building that is difficult to sustain alongside full-time product development\n'
       '\u2022 The monetization infrastructure (RevenueCat) is near-complete, which means the gap between \u201cbuilt\u201d and \u201cgenerating revenue\u201d is primarily a go-to-market execution problem, not a product problem\n\n'
       'We have not yet pitched investors because we want to have subscription revenue live and initial paying customers before entering fundraise conversations \u2014 which is why the grant and capital readiness support from Innovation Depot is timely.')

doc.add_page_break()

# ========== SECTION 3: SUPPORT PRIORITIES ==========
doc.add_heading('Support Priorities', level=1)
p = doc.add_paragraph('For each of the four Innovation Depot support programs, rate how much you need this type of support right now \u2014 on a scale of 1 (least needed) to 5 (most critical).')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'The Network \u2014 Priority Rating', '4')

add_qa(doc, 'The Network \u2014 Details',
       'We need targeted expertise in two areas: (1) Endurance sports industry partnerships \u2014 connections to running/triathlon coaching organizations, race event directors, and sports nutrition brands who could serve as pilot partners, early customers, or distribution channels. (2) Mobile app go-to-market strategy \u2014 operators who have scaled health/fitness subscription apps from pre-revenue to $50K+ MRR and can advise on pricing optimization, conversion funnel design, and user acquisition channels specific to the endurance athlete demographic.')

add_qa(doc, 'Special Projects Grant \u2014 Priority Rating', '5')

add_qa(doc, 'Special Projects Grant \u2014 Details',
       'We have previously received and successfully executed an Innovation Depot Special Projects Grant (funding Dr. Rachel Mitchell for sports nutrition algorithm validation), so we understand the format: a defined initiative, a specific vendor, and a measurable result.\n\n'
       'We are seeking $5,000\u201310,000 for a targeted go-to-market initiative to drive user acquisition and conversion ahead of our subscription launch. The initiative would include: (1) App Store Optimization and paid install campaigns on Apple Search Ads and Google Ads targeting endurance-related keywords (marathon nutrition, race fueling, triathlon nutrition plan) \u2014 estimated $3,000\u20135,000. Innovation Depot would pay Apple/Google directly. (2) Endurance influencer and content partnerships \u2014 sponsoring 2\u20133 running/triathlon content creators to demo the app and drive installs during peak race training season (spring/fall marathon cycles) \u2014 estimated $2,000\u20134,000. [CONFIRM: specific vendor/agency name if available]. The expected outcome is reaching 1,000+ downloads and establishing a measurable cost-per-install baseline that informs our fundraise materials.')

add_qa(doc, 'Capital Readiness \u2014 Priority Rating', '5')

add_qa(doc, 'Capital Readiness \u2014 Details',
       'We are 3\u20136 months from actively raising a pre-seed round of $250K\u2013$500K. We have not yet engaged investors or set a formal valuation. We need structured support in: (1) building a financial model that maps our dual revenue streams (B2C subscriptions + B2B coach subscriptions) to a defensible valuation, (2) refining our pitch deck and narrative for the sports nutrition technology space, (3) investor targeting \u2014 identifying angel investors and early-stage funds with health/fitness, sports tech, or consumer subscription thesis alignment, and (4) pitch readiness coaching. We have a completed pitch deck and product demo. This is our first raise beyond the $25K Alabama Launchpad non-dilutive prize, so structured guidance on round mechanics, term sheets, and investor diligence expectations would be highly valuable.')

add_qa(doc, 'Pilot Program \u2014 Priority Rating', '4')

add_qa(doc, 'Pilot Program \u2014 Details',
       'We have several ideal pilot scenarios, ordered by readiness:\n\n'
       '1. Coaching organizations: Partner with a running or triathlon coaching company (e.g., a local RRCA-affiliated coaching group or Birmingham Track Club coaching program) to offer Mealvana Endurance as a nutrition planning tool bundled with their coaching packages. Coach Mode is production-ready with web dashboard, athlete management, and messaging. We would scope a 3-month pilot where coaches use the platform to deliver nutrition guidance to 20\u201350 athletes, demonstrating retention and satisfaction metrics.\n\n'
       '2. Race event partners: Partner with a regional race organizer (e.g., Mercedes Marathon, Vulcan Run, or a Southeastern triathlon series) to offer Mealvana Endurance nutrition plans to registered participants as a value-add. This tests distribution through event partnerships and provides a large batch of engaged users.\n\n'
       '3. Sports nutrition brands: Partner with an endurance nutrition brand (e.g., a gel, hydration, or bar company) to feature their products within our AI-generated nutrition recommendations. This validates a product placement revenue stream and provides the brand with a data-driven customer acquisition channel.\n\n'
       '4. Sports dietitian practices: Partner with a sports dietitian practice to use Mealvana Endurance as their client-facing nutrition planning tool, replacing manual spreadsheet-based fueling plans.\n\n'
       'Our ICP for Coach Mode B2B is: certified endurance coaches (RRCA, USAT, ACE, NASM, ACSM) managing 10\u201350 athletes, currently using manual methods (spreadsheets, PDFs) for nutrition guidance, and charging $100\u2013$300/month per athlete for coaching packages.')

doc.add_page_break()

# ========== SECTION 4: BUSINESS & MARKET ==========
doc.add_heading('Business & Market', level=1)
p = doc.add_paragraph('The opportunity you are building into and the strength of your model')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'Problem You Are Solving',
       'Endurance athletes \u2014 runners, cyclists, and triathletes preparing for races from 5K to ultra-marathons \u2014 face a critical and underserved nutrition planning problem: they don\u2019t know what to eat, when to eat it, or how much to consume before, during, and after training and races.\n\n'
       'The cost of getting this wrong is significant:\n'
       '\u2022 GI distress during races (30\u201350% of marathon runners report it, per published ACSM data) caused by untested foods, improper timing, or excessive intake\n'
       '\u2022 Bonking / hitting the wall from insufficient carbohydrate fueling, which degrades performance by 10\u201320% in the final miles\n'
       '\u2022 Poor recovery that extends downtime between training sessions, increasing injury risk\n'
       '\u2022 Wasted money on race entry fees ($100\u2013$300+ per marathon), travel, and months of training undermined by a nutrition mistake on race day\n\n'
       'Today, athletes cobble together nutrition plans from generic blog posts, word-of-mouth, or trial-and-error. Coaches who want to provide nutrition guidance do so manually via spreadsheets and PDFs \u2014 time-consuming, inconsistent, and not personalized to the individual athlete\u2019s biometrics, food preferences, allergies, or gut tolerance.\n\n'
       'Mealvana Endurance solves this by generating personalized, science-based nutrition plans using ACSM metabolic formulas and ISSN sports nutrition guidelines, tailored to each athlete\u2019s body weight, pace, distance, dietary preferences, food allergies, gut training level, and even race-day weather conditions. We provide phase-specific guidance (pre-workout, during-workout, and recovery), recommend specific foods the athlete actually likes, and adjust hydration and sodium targets based on temperature and humidity.\n\n'
       'For coaches and sports dietitians, our Coach Mode eliminates manual nutrition planning entirely \u2014 they connect with athletes digitally, monitor and edit nutrition plans via a web dashboard, and deliver personalized guidance at scale.')

add_qa(doc, 'Market Size (TAM / SAM / SOM)',
       'TAM (Total Addressable Market): $4.2B\n'
       'The global sports nutrition market was valued at $45.2B in 2024 (Grand View Research), growing at 7.5% CAGR. Within this, the digital sports nutrition and fitness app segment represents approximately $4.2B, encompassing nutrition tracking, meal planning, and coaching platforms used by endurance and fitness athletes worldwide.\n\n'
       'SAM (Serviceable Addressable Market): $680M\n'
       'We focus specifically on endurance athletes (runners, cyclists, triathletes) in English-speaking markets (US, UK, Canada, Australia). There are approximately 60 million active runners in the US alone (Running USA), 47 million cyclists, and 4.4 million triathlon participants. Our SAM is the subset who race competitively or train for organized events \u2014 approximately 17 million in the US \u2014 and would pay for a specialized nutrition planning tool. At an average of $40/year (blended between casual and committed subscribers), this represents ~$680M.\n\n'
       'SOM (Serviceable Obtainable Market): $2M\u2013$5M in 5 years\n'
       'Our realistic 5-year capture target is 10,000\u201325,000 paying subscribers across B2C ($70/year) and B2B coach subscriptions ($240\u2013$600/year per coach). Bottom-up: if we reach 5,000 individual subscribers ($350K ARR) and 200 coaches averaging 15 athletes each ($120K ARR from coach subscriptions), we reach $470K ARR in year 3 and $2M\u2013$5M by year 5 through geographic expansion and deeper integration partnerships.\n\n'
       'Methodology: Bottom-up from race registration data (RunSignUp reports 10M+ US race registrations/year), training platform user bases (TrainingPeaks: 1M+ users; Strava: 120M+ users; FinalSurge: growing), and comparable health/fitness app conversion rates (1\u20133% free-to-paid).')

add_qa(doc, 'Market Trends & Timing',
       'Several converging trends make this the right time:\n\n'
       '1. Endurance sports participation is at record levels. Post-COVID, marathon finishers in the US rebounded to 650,000+ in 2023 (Running USA), and half-marathon participation exceeded 2 million. Triathlon participation is growing 8\u201310% annually. More participants means more demand for nutrition guidance.\n\n'
       '2. AI is transforming personalized health. The consumer expectation for AI-driven personalization has shifted dramatically in 2024\u20132026 \u2014 from generic recommendations to individually tailored plans. We ride this wave by using LLMs and linear programming to generate nutrition plans that are both scientifically rigorous and personally relevant.\n\n'
       '3. Endurance coaching is professionalizing and going digital. Platforms like TrainingPeaks, TrainerRoad, and FinalSurge have digitized training plans. Nutrition is the last major pillar of endurance performance that remains manual and fragmented. Our TrainingPeaks, FinalSurge, and Garmin integrations position us as the nutrition layer that complements existing training platforms.\n\n'
       '4. Sports nutrition product market is booming. The gels, bars, and hydration product market continues to grow with new brands entering constantly. These brands need data-driven channels to reach endurance athletes \u2014 our product recommendation engine creates a potential marketplace/affiliate revenue stream.\n\n'
       '5. Coach economy is expanding. More coaches are building independent practices and need digital tools to serve clients at scale. The shift to remote coaching (accelerated by COVID) means coaches need digital platforms for nutrition guidance, not in-person meal prep sessions.')

add_qa(doc, 'Competitive Landscape',
       'Direct competitors:\n\n'
       'Fuelin (fuelin.com): The closest competitor \u2014 AI-powered daily nutrition recommendations for endurance athletes, integrated with TrainingPeaks and Garmin. Strengths: established brand, pro athlete endorsements (used by Ironman athletes). Weaknesses: focused primarily on daily macros rather than race-specific fueling plans, higher price point (~$15/month), no coach mode, limited food preference personalization. We differentiate with deeper race-day specificity (phase-by-phase, hour-by-hour fueling plans), food preference learning, and our Coach Mode for B2B.\n\n'
       'Precision Fuel & Hydration (precisionfuelandhydration.com): Sweat testing and hydration planning for endurance athletes. Strengths: strong brand, partnerships with pro teams, product sales revenue. Weaknesses: primarily a product company (selling electrolytes) with a planning tool as a lead-gen feature, not a comprehensive nutrition planning platform. We compete on breadth \u2014 full pre/during/post nutrition plans, not just hydration.\n\n'
       'Hexis (hexisperformance.com): AI daily nutrition planner for athletes. Strengths: clean UX, periodized nutrition. Weaknesses: broader athlete focus (not endurance-specific), newer to market, limited integrations.\n\n'
       'Indirect competitors:\n\n'
       'MyFitnessPal / Cronometer: General calorie and macro tracking apps. Not endurance-specific, no race fueling plans, no AI-generated recommendations tailored to activity parameters. We do not compete on daily food logging \u2014 we solve the race-day and training-day fueling problem specifically.\n\n'
       'TrainingPeaks / Strava: Training platforms, not nutrition platforms. They are integration partners, not competitors. Neither offers nutrition plan generation.\n\n'
       'Our sustainable competitive advantages:\n'
       '1. Dual AI + algorithmic system: Our LLM-powered planning with deterministic fallback ensures both quality and reliability \u2014 no competitor offers this hybrid approach\n'
       '2. Food preference learning: Three-tier preference system (love/willing-to-try/avoid) that personalizes recommendations in a way no competitor matches\n'
       '3. Coach Mode (B2B): Built-in many-to-many coaching platform with web dashboard \u2014 Fuelin and Hexis have no equivalent\n'
       '4. Integration depth: Live integrations with TrainingPeaks, FinalSurge, and Garmin. Our TrainingPeaks write-back (pushing nutrition data back to the training platform) is unique\n'
       '5. Offline-first architecture: Full functionality without internet \u2014 critical for athletes in remote training locations\n'
       '6. Evidence-based foundation: Built on validated ACSM metabolic formulas and ISSN guidelines, not generic nutrition advice\n\n'
       'We acknowledge that Fuelin has a head start on brand awareness and athlete endorsements. Our advantage is product depth, coach integration, and the ability to serve both individual athletes and coaching organizations as a platform.')

add_qa(doc, 'Business Model',
       'Dual revenue model: B2C subscriptions + B2B coach subscriptions.\n\n'
       'B2C (Individual Athletes):\n'
       '\u2022 Free tier: Basic nutrition plan generation, food preference tracking, gut training guidance\n'
       '\u2022 Pro tier: $9.99/month or $69.99/year (1-month free trial)\n'
       '\u2022 Pro features: Training platform integrations (TrainingPeaks, FinalSurge, Garmin write-back), Coach Mode access, brick workout nutrition, personal fueling templates, Mealvana 101 video course, by-hour race-day planning, food journaling, barcode scanning, carb-loading protocols, adaptive macro adjustment\n'
       '\u2022 Payment: RevenueCat (in-app subscriptions via Apple/Google)\n\n'
       'B2B (Coaches & Dietitians):\n'
       '\u2022 Free during beta (current phase)\n'
       '\u2022 Starter: $20/month (up to 20 athletes)\n'
       '\u2022 Pro: $50/month (unlimited athletes, advanced analytics)\n'
       '\u2022 Payment: Stripe (planned)\n'
       '\u2022 Sales motion: Direct outreach to RRCA, USAT, ACE, NASM, ACSM-certified coaches; partnerships with coaching organizations; content marketing targeting coaches\n\n'
       'Future revenue streams (12\u201324 month horizon):\n'
       '\u2022 Product placement / affiliate revenue from sports nutrition brands featured in recommendations\n'
       '\u2022 Enterprise partnerships with race event organizers (bulk athlete access)\n'
       '\u2022 API licensing for training platforms seeking a nutrition intelligence layer')

add_qa(doc, 'Revenue & Traction',
       'We are currently pre-revenue. Our most significant traction signals:\n\n'
       '\u2022 100\u2013500 app users across iOS and Android with organic discovery (no paid acquisition to date)\n'
       '\u2022 18 product versions shipped (v1.0 through v1.18.1) since launch, demonstrating rapid iteration and product maturity\n'
       '\u2022 Positive beta user feedback driving feature development \u2014 users have requested and validated features like carb loading, food preference learning, and coach mode\n'
       '\u2022 Active coach interest from certified coaches and sports dietitians seeking a digital nutrition planning tool\n'
       '\u2022 Production-ready Coach Mode v2 with web dashboard, ready for onboarding\n'
       '\u2022 Live integrations with TrainingPeaks (credentials approved December 2025), FinalSurge, and Garmin\n'
       '\u2022 Mixpanel analytics active and tracking user behavior, establishing baseline engagement data\n'
       '\u2022 $25,000 Alabama Launchpad prize (2021) as external validation of concept and team\n'
       '\u2022 Full monetization infrastructure nearly complete (RevenueCat implementation in progress, 11 Pro features identified and gated)\n\n'
       'We expect to have subscription revenue live within 60 days of this application.')

add_qa(doc, 'Regulatory Environment',
       'Our market is not heavily regulated. Mealvana Endurance provides nutrition planning guidance, not medical advice, and we make this distinction explicit in our Terms of Service and in-app disclaimers. Key considerations:\n\n'
       '\u2022 We do not diagnose, treat, or prescribe \u2014 our plans are informational and evidence-based, similar to content from running magazines or coaching websites\n'
       '\u2022 Our Coach Mode verifies professional certifications (RRCA, USAT, ACE, NASM, ACSM) and professional liability insurance, which demonstrates responsible platform design\n'
       '\u2022 App Store health claims policies require careful language \u2014 we frame our science-based recommendations as \u201cguidance\u201d rather than \u201ctreatment\u201d\n'
       '\u2022 We collect minimal personal data (device-based authentication is default, no user accounts required) and are transparent about data practices in our privacy policy\n'
       '\u2022 Dr. Rachel Mitchell, our part-time sports dietitian, reviews nutrition algorithm outputs for clinical appropriateness',
       required=False)

doc.add_page_break()

# ========== SECTION 5: TEAM & FOUNDER ==========
doc.add_heading('Team & Founder', level=1)
p = doc.add_paragraph('Who you are, what makes you uniquely positioned, and how you lead')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'Founder Background',
       'Xuan Huang \u2014 CEO & Co-founder\n'
       'Xuan is a statistician by training and an engineer at heart. Her background in data science and analytics is the foundation of Mealvana\u2019s personalized nutrition algorithms. She personally experienced the meal planning problem \u2014 first as a busy professional trying to answer \u201cwhat\u2019s for dinner?\u201d (which inspired the original Mealvana consumer app) and then watching endurance athlete friends struggle with race nutrition. Before Milkman, she built data products in professional settings and has deep expertise in statistical modeling, machine learning, and recommendation systems. She has prior startup experience and has built industry connections in the Birmingham startup ecosystem through Innovation Depot, Forge, and Alabama Launchpad.\n\n'
       'Lee Martin \u2014 CTO & Co-founder\n'
       'Lee is a full-stack engineer and competitive triathlete whose personal experience with race nutrition directly shaped the product. He has extensive experience in Flutter mobile development, Supabase backend architecture, and AI/ML integration. His meticulous approach to design and user experience \u2014 combined with firsthand understanding of what endurance athletes need during training and races \u2014 makes him uniquely positioned to build the technical product. Lee architected the dual AI + algorithmic nutrition planning system, the offline-first architecture, and the Coach Mode web portal.\n\n'
       'Together, Xuan and Lee combine data science rigor with engineering execution and genuine domain expertise as athletes and food enthusiasts. They are supported by Dr. Rachel Mitchell (sports dietitian providing clinical validation), Rui Tian (user researcher), and Haibin Zhuang (engineering consultant).')

add_qa(doc, 'LinkedIn Profile(s)',
       'Xuan Huang (Founder/CEO): https://www.linkedin.com/in/xuan-huang/\n'
       'Lee Martin (CTO): https://www.linkedin.com/in/lee-martin-developer')

add_qa(doc, 'Team & Next Hire',
       'Current core team:\n'
       '\u2022 Xuan Huang \u2014 CEO & Co-founder (full-time). Strategy, data science, business development, product direction\n'
       '\u2022 Lee Martin \u2014 CTO & Co-founder (full-time). Engineering, architecture, product development, design\n'
       '\u2022 Dr. Rachel Mitchell \u2014 Sports Dietitian (part-time). Clinical validation of nutrition algorithms, Mealvana 101 educational content\n'
       '\u2022 Rui Tian \u2014 User Researcher (part-time). User interviews, feedback synthesis, UX research\n'
       '\u2022 Haibin Zhuang \u2014 Engineering Consultant. Backend infrastructure and scalability\n\n'
       'Next hire: Head of Growth / Sales & Partnerships\n\n'
       'This is our most critical gap. Both founders are deep in product \u2014 which has produced a robust, feature-rich app \u2014 but we need a dedicated person driving user acquisition, coach outreach, and partnership development. The ideal hire is someone with experience in health/fitness app marketing, B2B SaaS sales (specifically selling to coaches or fitness professionals), and endurance sports community connections. This hire directly addresses our core constraint: converting a strong product into revenue requires dedicated go-to-market execution, not more engineering.')

add_qa(doc, 'Long-Term Vision', 'Build to scale')

add_qa(doc, 'Long-Term Vision \u2014 Explanation',
       'Our long-term vision is to become the nutrition intelligence layer for the entire endurance sports ecosystem.\n\n'
       'Today, training platforms (TrainingPeaks, Strava, Garmin) tell athletes how to train. Mealvana tells them how to fuel. We believe nutrition planning will become as essential and integrated as training plans \u2014 and that the company that owns the intelligent, personalized nutrition data layer will capture significant value.\n\n'
       'The endgame in milestones:\n\n'
       '1. Near-term (12 months): Prove B2C and B2B product-market fit with paying subscribers and active coaches. Close pre-seed round.\n\n'
       '2. Mid-term (2\u20133 years): Expand integrations to become the default nutrition engine within the major training platforms (TrainingPeaks, Strava, Garmin, Wahoo). Build a product recommendation marketplace where nutrition brands pay for intelligent placement in athlete fueling plans. Reach $2M+ ARR.\n\n'
       '3. Long-term (5+ years): Become the platform where every endurance athlete gets their personalized nutrition plan \u2014 whether directly through our app, through their coach, through their training platform, or through their race event. Expand from endurance sports into broader sports nutrition (team sports, strength athletes, recreational fitness). License our nutrition intelligence engine to enterprise partners (hospital wellness programs, corporate fitness, military performance optimization).\n\n'
       '4. Scale milestone: 100,000+ athletes using Mealvana-powered nutrition plans (directly or through partners), generating $10M+ ARR through a mix of consumer subscriptions, coach/enterprise subscriptions, and marketplace/API revenue.\n\n'
       'We are building a standalone, venture-scale nutrition technology company \u2014 not a lifestyle business. The endurance sports market is large enough and underserved enough in nutrition specifically to support this trajectory, and our technical architecture (AI + evidence-based algorithms + offline-first + coach platform + integration layer) is purpose-built for platform scale.')

doc.add_page_break()

# ========== SECTION 6: PRODUCT & BRAND ==========
doc.add_heading('Product & Brand', level=1)
p = doc.add_paragraph('Where your product and brand stand today \u2014 and where they are headed')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True
add_divider(doc)

add_qa(doc, 'Product Demo',
       '[TODO: Upload Loom video \u2014 5\u201310 minutes walking through the app: onboarding, plan generation, food preferences, Coach Mode, integrations, carb loading]')

add_qa(doc, 'Product Roadmap',
       'Our top 3 product priorities for the next 6\u201312 months:\n\n'
       '1. Subscription monetization via RevenueCat (Q2 2026)\n'
       'What: Launch the paid Pro tier ($9.99/month, $69.99/year) gating 11 premium features behind a paywall.\n'
       'Why: Beta user feedback confirms willingness to pay for training platform integrations, coach access, and advanced fueling tools. We cannot sustain development on bootstrapped capital alone. Every conversation with potential investors starts with \u201cdo you have revenue?\u201d \u2014 this is the prerequisite for our fundraise.\n'
       'Driving signal: Multiple beta users have asked when paid features will be available; coach interest is contingent on a sustainable business model behind the platform.\n\n'
       '2. Coach Mode go-to-market and onboarding optimization (Q2\u2013Q3 2026)\n'
       'What: Streamline coach registration, automate certification verification (currently manual), build coach-facing marketing materials, and onboard first 10+ paying coaches.\n'
       'Why: Coach Mode v2 is production-ready but untested at scale. User research (Rui Tian) indicates coaches want a nutrition tool that integrates with their existing workflow. The B2B channel has higher LTV per customer and provides a built-in distribution mechanism (each coach brings 10\u201350 athletes).\n'
       'Driving signal: Direct conversations with RRCA-certified coaches who currently deliver nutrition guidance via manually created PDFs and spreadsheets.\n\n'
       '3. Deeper Garmin and wearable data integration (Q3\u2013Q4 2026)\n'
       'What: Use Garmin push data (heart rate, sleep, stress, body composition) to auto-adjust nutrition recommendations based on recovery status, training load, and biometric trends.\n'
       'Why: The integration infrastructure is built (Garmin push endpoints live, edge functions deployed), but we are currently only matching completed activities to planned ones. Closing the loop \u2014 using wearable data to personalize nutrition proactively \u2014 is a significant competitive differentiator that no competitor currently offers.\n'
       'Driving signal: Beta user requests for \u201csmarter\u201d recommendations that account for how they slept, their HRV, and their training load trends.')

add_qa(doc, 'Intellectual Property',
       'We have filed trademarks for \u201cMilkman\u201d and \u201cMealvana.\u201d Our primary proprietary technology includes:\n\n'
       '\u2022 Dual AI + algorithmic nutrition planning engine: A hybrid system combining LLM-powered planning with a deterministic fallback algorithm using ACSM metabolic formulas, ISSN guidelines, and linear programming optimization. This dual-system approach (AI-first with algorithmic reliability guarantee) is architecturally unique in the sports nutrition space.\n'
       '\u2022 Food preference learning system: A three-tier preference model (love/willing-to-try/avoid) with phase-specific food scoring that personalizes recommendations while maintaining nutritional targets \u2014 a proprietary approach to balancing athlete preferences with science-based requirements.\n'
       '\u2022 \u201cFat backend\u201d content management system: A server-side architecture that allows nutrition algorithm parameters, food databases, and UI content to be updated without app store releases \u2014 enabling rapid A/B testing and continuous improvement of nutrition recommendations.\n\n'
       'We are evaluating whether to file a provisional patent on the dual-system nutrition planning methodology.',
       required=False)

add_qa(doc, 'Customer Validation',
       'Our best validation evidence:\n\n'
       '\u2022 100\u2013500 organic users with zero paid acquisition spend \u2014 all growth through word-of-mouth and organic app store discovery, indicating genuine demand\n'
       '\u2022 Positive beta feedback driving product development: users specifically requested and validated carb loading protocols, food preference learning, and Coach Mode \u2014 features we built in direct response to user input\n'
       '\u2022 Coach and dietitian interest: Sports dietitians and certified coaches have expressed interest in using the platform as a client-facing tool, validating the B2B opportunity. Dr. Rachel Mitchell (sports dietitian) joined the team part-time after seeing the product, which is itself a strong validation signal.\n'
       '\u2022 18 shipped versions incorporating user feedback, demonstrating product-market dialogue (not just building in a vacuum)\n'
       '\u2022 TrainingPeaks credential approval (December 2025): TrainingPeaks reviewed our product and granted API access with broad scopes (athlete profile, events, nutrition, workouts, metrics, webhooks) \u2014 third-party platform validation of our product\u2019s quality and value proposition\n'
       '\u2022 Alabama Launchpad prize ($25K, 2021): External panel of judges and investors validated the concept and team at the competition stage')

add_qa(doc, 'Pitch Deck', '[TODO: Upload latest pitch deck here]')

add_qa(doc, 'Additional Materials', '[TODO: Upload any one-pagers, press coverage, or supporting materials]', required=False)

add_qa(doc, 'Submission Date', '04/08/2026')

# ========== SAVE ==========
output_path = os.path.expanduser('~/development/mealvana_endurance/docs/_archived/business/Innovation_Depot_Application_Mealvana.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
